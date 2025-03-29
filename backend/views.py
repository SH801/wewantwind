import os
from docx import Document
import sys
import docx
import json
import requests
import shutil
import uuid
import urllib
import geojson
import subprocess
import psycopg2
from psycopg2.extensions import AsIs
from turfpy.transformation import intersect
from turfpy.measurement import points_within_polygon
from geojson import Point as GeoJSONPoint, Feature, FeatureCollection, Polygon
from urllib.parse import urlparse
from PIL import Image, ImageEnhance
from matplotlib import colors
from pprint import pprint
from ipware import get_client_ip
from docx.shared import Cm, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from reportlab.pdfgen.canvas import Canvas
from reportlab.lib.units import cm, inch
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from django.http import HttpResponse, HttpResponseNotFound
from django.shortcuts import render
from django.views.decorators.csrf import csrf_exempt
from django.contrib.gis.geos import Point
from django.contrib.gis.db.models.functions import Distance
from django.contrib.gis.measure import Distance as RadiusDistance
from django.http import HttpResponse
from django.core.serializers.json import DjangoJSONEncoder
from django.template.loader import render_to_string, get_template
from django.core.mail import EmailMessage
from django.contrib.gis.geos import Point
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from django.contrib.sites.shortcuts import get_current_site
from django.utils.encoding import force_bytes
from django.shortcuts import redirect
from django.db.models import Q, Count
from django.http import HttpResponse
from django.utils.encoding import iri_to_uri
from time import sleep
from random import randrange
from itertools import islice
from osgeo import gdal, osr, ogr
import numpy as np
import json
from turfpy.misc import line_arc
from geojson import Feature as GeoJSONFeature, Point as GeoJSONPoint

gdal.DontUseExceptions()

from .models import Site, Vote, Message, Boundary, EventLog

# Create your views here.

cwd = os.path.dirname(os.path.realpath(__file__))

COORDINATE_PRECISION = 5
LOCAL_DISTANCE = 10 # miles

RESTART_SCRIPT = os.environ.get("RESTART_SCRIPT")
GOOGLE_RECAPTCHA_SECRET_KEY = os.environ.get("GOOGLE_RECAPTCHA_SECRET_KEY")

VIEWSHED_MAX_CIRCULAR_RANGE = float(45000)
VIEWSHED_MAX_DISTANCE = float((2 * (VIEWSHED_MAX_CIRCULAR_RANGE ** 2)) ** 0.5)

TERRAIN_FILE = cwd + '/terrain/terrain_lowres_withfeatures.tif'

DEFAULT_HUB_HEIGHT = 108.3
DEFAULT_BLADE_RADIUS = 69.25

constraintslist = [
    {
        'heading': 'All constraints',
        'datasets': [
            'Inadequate wind speed',
            'Landscape and visual impact',
            'Heritage impacts',
            'Separation distance to residential properties',
            'Ecology and wildlife',
            'Other technical constraints',
            'Aviation and exclusion areas'
        ],
        'layers': [
            {'name': 'wind', 'color': 'blue'},
            {'name': 'landscape_and_visual_impact', 'color': 'chartreuse'},
            {'name': 'heritage_impacts', 'color': 'darkgoldenrod'},
            {'name': 'separation_distance_to_residential', 'color': 'darkorange'},
            {'name': 'ecology_and_wildlife', 'color': 'darkgreen'},
            {'name': 'other_technical_constraints_hi', 'color': 'red'},
            {'name': 'aviation_and_exclusion_areas', 'color': 'purple'}
        ]
    },
    {
        'heading': 'Inadequate wind speed', 
        'datasets': [
            'Inadequate wind speed (< 5 metres per second) from NOABL wind speed database'
        ],
        'layers': [
            {'name': 'wind', 'color': 'blue'}
        ]
    },
    {
        'heading': 'Landscape and visual impact', 
        'datasets': [
            'National Parks', 
            'Areas of Outstanding Natural Beauty / National Scenic Areas', 
            'Heritage Coasts'
        ],
        'layers': [
            {'name': 'landscape_and_visual_impact', 'color': 'chartreuse'}
        ]
    },
    {
        'heading': 'Heritage impacts', 
        'datasets': [
            'Listed buildings',
            'Conservation areas',
            'World Heritage Sites',
            'Scheduled Ancient Monuments',
            'Registered parks and gardens',
            'Registered historic battlefields'
        ],
        'layers': [
            {'name': 'heritage_impacts', 'color': 'darkgoldenrod'}
        ]
    },
    {
        'heading': 'Separation distance to residential properties', 
        'datasets': [
            '400 metre buffer from residential areas'
        ],
        'layers': [
            {'name': 'separation_distance_to_residential', 'color': 'darkorange'}
        ]
    },
    {
        'heading': 'Ecology and wildlife', 
        'datasets': [
            'Sites of Special Scientific Interest / Areas of Special Scientific Interest',
            'Ramsar sites',
            'Special Areas of Conservation',
            'Special Protection Areas',
            'National nature reserves',
            'Ancient woodland',
            'Local wildlife reserves',
            '50 metre buffer from hedgerows'
        ],
        'layers': [
            {'name': 'ecology_and_wildlife', 'color': 'darkgreen'}
        ]
    },
    {
        'heading': 'Other technical constraints', 
        'datasets': [
            '150 metre buffer from public roads (A and B roads and motorways)',
            '150 metre buffer from railway lines',
            '150 metre buffer from inland waters',
            '150 metre buffer from pipelines',
            '150 metre buffer from power lines',
            '150 metre buffer from public footpaths',
            '200 metre buffer from bridleways',     
        ],
        'layers': [
            {'name': 'other_technical_constraints_hi', 'color': 'red'}
        ]
    },
    {
        'heading': 'Aviation and exclusion areas', 
        'datasets': [
            'Civilian airports',
            'Aerodromes',
            'Military airfields and airbases',
            'MOD training areas',
            'Explosive safeguarded areas, danger areas near ranges',
            'MOD exclusion areas'
        ],
        'layers': [
            {'name': 'aviation_and_exclusion_areas', 'color': 'purple'}
        ]
    }
]


def OutputJson(json_array={'result': 'failure'}):
    json_data = json.dumps(json_array, cls=DjangoJSONEncoder, indent=2)
    return HttpResponse(json_data, content_type="text/json")

def OutputError():
    return OutputJson()

def SendErrorEmail(messagecontent):
    """
    Send message to sysadmin in event of error
    """
    message = EmailMessage("WeWantWind.org - Error message", messagecontent, from_email="info@wewantwind.org", to=[os.environ.get("ERROR_EMAIL")])
    message.send()

def home(request):
    """
    Shows default home page or other frontend-specific pages to be rendered by frontend React app
    """
    return render(request, 'index.html')

def GetRandomPointInBounds():
    bounds = [
        [
            -4.9,
            50
        ],
        [
            1.5, 
            57 
        ]
    ]

    xmin, ymin, xmax, ymax = bounds[0][0], bounds[0][1], bounds[1][0], bounds[1][1]
    xrandom = float(randrange(int(xmin * 100), int(xmax * 100), 1) / 100)
    yrandom = float(randrange(int(ymin * 100), int(ymax * 100), 1) / 100)
    return [xrandom, yrandom]

def CheckInBoundary(point):
    """
    Checks whether point within UK using Boundary table which holds UK
    """
    point = Point(float(point[0]), float(point[1]))
    return (Boundary.objects.filter(name='UK Boundary').filter(geometry__intersects=point).count() != 0)

@csrf_exempt
def GetRandomPoint(request):
    """
    Get random point within UK
    """

    while True:
        point = GetRandomPointInBounds()
        inboundary = CheckInBoundary(point)
        if inboundary is False:
            pass 
            # print("Point not within UK", point)
        else: return OutputJson({'lng': point[0], 'lat': point[1]})

def postgisGetResultsAsDict(sql_text, sql_parameters=None):
    """
    Runs database query and returns results
    """

    POSTGRES_DB = os.environ['POSTGRES_DB']
    POSTGRES_USER = os.environ['POSTGRES_USER']
    POSTGRES_PASSWORD = os.environ['POSTGRES_PASSWORD']

    conn = psycopg2.connect(dbname=POSTGRES_DB, user=POSTGRES_USER, password=POSTGRES_PASSWORD)
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    if sql_parameters is None: cur.execute(sql_text)
    else: cur.execute(sql_text, sql_parameters)
    results = cur.fetchall()
    conn.close()
    return results

def GetNearestTurbine(lat, lng):
    """
    Get nearest optimal wind turbine site to specified point
    """

    sites = postgisGetResultsAsDict("""  
    SELECT 
    id,
    ST_X(centre) turbinelng,
    ST_Y(centre) turbinelat,
    ST_Distance(ST_Transform(ST_SetSRID(ST_MakePoint(%s, %s), 4326), 3857), ST_Transform(centre, 3857)) distance, 
    ST_Area(ST_Transform(geometry, 3857)) area 
    FROM backend_site 
    ORDER BY distance LIMIT 50;
    """, (AsIs(lng), AsIs(lat), ))
    if len(sites) == 0: return None

    largestarea, largestindex, num_to_check = None, 0, 5
    sites = sites[0:num_to_check]
    for i in range(num_to_check):
        site = sites[i]
        area_square_meters = site['area']
        if ((largestarea is None) or (area_square_meters > largestarea)):
            largestindex = i
            largestarea = area_square_meters

    site = sites[largestindex]

    site['currentlng'] = lng
    site['currentlat'] = lat
    site['turbinelng'] = site['turbinelng']
    site['turbinelat'] = site['turbinelat']
    site['distance_mi'] = (site['distance'] * 0.621371) / 1000
    site['distance_km'] = site['distance'] / 1000
    site['distance_m'] = site['distance']

    return site

# *** Potentially broken install of GeoDjango

def GetNearestTurbine_Broken(lat, lng):
    """
    Get nearest optimal wind turbine site to specified point
    """

    centre = Point(lng, lat, srid=4326)    
    firstcutsize = 50
    sites = Site.objects.all().annotate(centredistance=Distance('centre' , centre )).order_by('centredistance')[:firstcutsize].annotate(distance=Distance('geometry' , centre))
    if sites is None: return None

    firstcut = {}
    for site in sites: firstcut[site.distance] = site
    sortedfirstcut = dict(sorted(firstcut.items()))

    largestarea = None
    largestindex = 0
    num_to_check = 5
    sites = list(islice(sortedfirstcut.items(), num_to_check))
    # print(sites)
    for i in range(num_to_check):
        site = sites[i][1]
        area_square_meters = site.geometry.area
        # print(i, site.centre, int(100000 * area_square_meters))
        # print(area_square_meters)
        if ((largestarea is None) or (area_square_meters > largestarea)):
            largestindex = i
            largestarea = area_square_meters

    # print("largestindex", largestindex)
    site = sites[largestindex][1]
    return site

@csrf_exempt
def NearestTurbine(request):
    """
    Get nearest turbine to supplied point
    """

    try:
        latlng = json.loads(request.body)
    except ValueError:
        return OutputError()

    # lng, lat = -0.147562, 50.8289154

    site = GetNearestTurbine(latlng['lat'], latlng['lng'])
    if site is None: return OutputError()

#    results = {}
#    results['currentlng'] = latlng['lng']
#    results['currentlat'] = latlng['lat']
#    results['turbinelng'] = site.centre.coords[0]
#    results['turbinelat'] = site.centre.coords[1]
#    results['distance_mi'] = site.distance.mi
#    results['distance_km'] = site.distance.km
#    results['distance_m'] = site.distance.m

    results = site

    ip, is_routable = get_client_ip(request)
    eventlog = EventLog(name='NearestTurbine', content=json.dumps(results, indent=2), ip=ip)
    eventlog.save()

    return OutputJson(results)

@csrf_exempt
def CheckRenderer(request):
    if request.user.is_authenticated is False:
        return OutputJson({'error': 'You need to be logged in'})
    while True:
        coordinates = GetRandomPointInBounds()
        inboundary = CheckInBoundary(coordinates)
        if inboundary is True: break

    threedimensionsparameters = {'width': '600', 'height': '600', 'ratio': '3', 'zoom': '15', 'pitch': '45', 'bearing': '0', 'center': str(coordinates[0]) + ',' + str(coordinates[1])}
    with open(cwd + '/styles/3d.json', encoding='utf-8') as fp: threedimensions = json.load(fp)
    threedimensions['style']['sources']['customgeojson'] = {
        "type": "geojson",
        "data": {"type": "FeatureCollection", "features": [{
            "type": "Feature",
            "geometry": {
                "type": "Point",
                "coordinates": coordinates
            }
    }]}}
    for key in threedimensionsparameters: threedimensions[key] = threedimensionsparameters[key]
    try:
        r = requests.post('http://localhost:81/render', json=threedimensions, stream=True)
        if r.status_code == 200:
            return HttpResponse(r.raw, content_type="image/png")        
        else:
            return OutputJson(r)
    except requests.exceptions.ConnectionError:
        return OutputJson({'error': 'Unable to connect to tilerenderer server'})

@csrf_exempt
def CheckRestartRenderer(request):
    while True:
        coordinates = GetRandomPointInBounds()
        inboundary = CheckInBoundary(coordinates)
        if inboundary is True: break
    threedimensionsparameters = {'width': '5', 'height': '5', 'ratio': '1', 'zoom': '15', 'pitch': '45', 'bearing': '0', 'center': str(coordinates[0]) + ',' + str(coordinates[1])}
    with open(cwd + '/styles/3d.json', encoding='utf-8') as fp: threedimensions = json.load(fp)
    threedimensions['style']['sources']['customgeojson'] = {
        "type": "geojson",
        "data": {"type": "FeatureCollection", "features": [{
            "type": "Feature",
            "geometry": {
                "type": "Point",
                "coordinates": coordinates
            }
    }]}}
    for key in threedimensionsparameters: threedimensions[key] = threedimensionsparameters[key]
    try:
        r = requests.post('http://localhost:81/render', json=threedimensions, stream=True)
        if r.status_code == 200:
            return HttpResponse(r.raw, content_type="image/png")        
        else:
            return OutputJson(r)
    except requests.exceptions.ConnectionError:
        SendErrorEmail("Failed to load tilerenderer so restarting")
        sys.stdout.reconfigure(encoding='utf-8')
        shelloutput = subprocess.run(RESTART_SCRIPT, encoding="utf8", capture_output=True, text=True, universal_newlines=True) 
        return OutputJson({'result': shelloutput.stderr})

@csrf_exempt
def RestartRenderer(request):
    if request.user.is_authenticated is False:
        return OutputJson({'error': 'You need to be logged in'})

    sys.stdout.reconfigure(encoding='utf-8')
    shelloutput = subprocess.run(RESTART_SCRIPT, encoding="utf8", capture_output=True, text=True, universal_newlines=True) 
    return OutputJson({'result': shelloutput.stderr})

def processimages(id, coordinates, constraintslist, parameters, hubheight, bladeradius):
    with open(cwd + '/styles/planningconstraint.json', encoding='utf-8') as fp: planningconstraint = json.load(fp)
    with open(cwd + '/styles/planningconstraints_osmstyle.json', encoding='utf-8') as fp: planningconstraints = json.load(fp)
    with open(cwd + '/styles/viewshed_osmstyle.json', encoding='utf-8') as fp: viewshed = json.load(fp)
    # with open(cwd + '/styles/planningconstraints.json', encoding='utf-8') as fp: planningconstraints = json.load(fp)

    imagedirectory = cwd + '/downloads/' + id
    os.mkdir(imagedirectory) 

    threedimensionsparameters = {'width': '600', 'height': '600', 'ratio': '3', 'zoom': '15', 'pitch': '45', 'bearing': '0', 'center': str(coordinates[0]) + ',' + str(coordinates[1])}

    with open(cwd + '/styles/3d.json', encoding='utf-8') as fp: threedimensions = json.load(fp)
    threedimensions['style']['sources']['customgeojson'] = {
        "type": "geojson",
        "data": {"type": "FeatureCollection", "features": [{
            "type": "Feature",
            "geometry": {
                "type": "Point",
                "coordinates": coordinates
            }
    }]}}
    for key in threedimensionsparameters: threedimensions[key] = threedimensionsparameters[key]

    try:
        r = requests.post('http://localhost:81/render', json=threedimensions, stream=True)
        if r.status_code == 200:
            imagepath = imagedirectory + '/3d.png'
            with open(imagepath, 'wb') as f:
                r.raw.decode_content = True
                shutil.copyfileobj(r.raw, f)
    except requests.exceptions.ConnectionError:
        SendErrorEmail("Couldn't connect to tilerenderer server")

    windturbine = Image.open(cwd + "/windturbine_bright.png")
    windturbinescale = 0.4
    windturbine = windturbine.resize((int(windturbinescale * windturbine.size[0]), int(windturbinescale * windturbine.size[1])), Image.Resampling.LANCZOS)
    background = Image.open(imagepath)
    contrast = ImageEnhance.Contrast(background)
    brightness = ImageEnhance.Brightness(contrast.enhance(0.8))
    brightness.enhance(1.2).save(imagepath)    
    background = Image.open(imagepath)
    background.paste(windturbine, (int((int(threedimensionsparameters['width']) * 3 / 2) - (windturbine.size[0] / 2)), int(int(threedimensionsparameters['height']) * 3 / 2) - windturbine.size[1]), windturbine)

    background.save(imagepath)

    viewshedparameters = {'width': '600', 'height': '600', 'ratio': '3', 'zoom': '8.3', 'pitch': '0', 'bearing': '0', 'center': str(coordinates[0]) + ',' + str(coordinates[1])}
    viewshedgeojson = GetViewsheds(coordinates[0], coordinates[1], hubheight, bladeradius)

    if 'viewshed' in viewshed['style']['sources']:
        viewshed['style']['sources']['viewshed']['data'] = viewshedgeojson

    if 'customgeojson' in viewshed['style']['sources']:
        viewshed['style']['sources']['customgeojson'] = {
            "type": "geojson",
            "data": {"type": "FeatureCollection", "features": [{
                "type": "Feature",
                "geometry": {
                    "type": "Point",
                    "coordinates": coordinates
                }
            }]}
        }      

    for key in viewshedparameters: viewshed[key] = viewshedparameters[key]

    try:
        r = requests.post('http://localhost:81/render', json=viewshed, stream=True)
        if r.status_code == 200:
            imagepath = imagedirectory + '/viewshed.png'
            with open(imagepath, 'wb') as f:
                r.raw.decode_content = True
                shutil.copyfileobj(r.raw, f)
    except requests.exceptions.ConnectionError:
        SendErrorEmail("Couldn't connect to tilerenderer server")

    # Modify planningconstraint template according to constraint color and duplicating for each listed layer
    for constraint in constraintslist:
        # print(constraint['heading'])
        alllayers = []
        constraintlayers = json.loads(json.dumps(planningconstraint))
        for layer in constraint['layers']:
            for constraintlayer in constraintlayers:
                newconstraintlayer = json.loads(json.dumps(constraintlayer))
                newconstraintlayer['id'] = layer['name'] + "_" + newconstraintlayer['id']
                newconstraintlayer['source-layer'] = layer['name']
                if 'fill-color' in newconstraintlayer['paint']: newconstraintlayer['paint']['fill-color'] = layer['color']
                alllayers.append(newconstraintlayer)

        # Insert new layers into existing stylesheet template
        newlayers = []
        newplanningconstraints = json.loads(json.dumps(planningconstraints))
        for newplanningconstraintlayer in newplanningconstraints['style']['layers']:
            if newplanningconstraintlayer['id'] == 'planningconstraints':
                for layer in alllayers:
                    newlayers.append(layer)
            else: 
                newlayers.append(newplanningconstraintlayer)

        if 'customgeojson' in newplanningconstraints['style']['sources']:
            newplanningconstraints['style']['sources']['customgeojson'] = {
                "type": "geojson",
                "data": {"type": "FeatureCollection", "features": [{
                    "type": "Feature",
                    "geometry": {
                        "type": "Point",
                        "coordinates": coordinates
                    }
                }]}
            }      

        newplanningconstraints['style']['layers'] = newlayers
        for key in parameters: newplanningconstraints[key] = parameters[key]

        foreground = Image.open(cwd + "/scale.png")
        # For some reason the meticulously sized scale doesn't seem to fit to the correct image size
        # Possibly some quirk with the map renderer so need to scale 'by hand'!
        scalescale = 1.5
        foreground = foreground.resize((int(scalescale * foreground.size[0]), int(scalescale * foreground.size[1])), Image.Resampling.LANCZOS)
        # pprint(newplanningconstraints, indent=4)
        r = requests.post('http://localhost:81/render', json=newplanningconstraints, stream=True)
        if r.status_code == 200:
            imagepath = imagedirectory + "/" + constraint['heading'] + '.png'
            with open(imagepath, 'wb') as f:
                r.raw.decode_content = True
                shutil.copyfileobj(r.raw, f) 
            background = Image.open(imagepath)
            background.paste(foreground, ((600 * 3) - foreground.size[0], (500 * 3) - foreground.size[1]), foreground)
            background.save(imagedirectory + "/" + constraint['heading'] + '.png')

    return imagedirectory

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(
        docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    new_run.style = get_or_create_hyperlink_style(part.document)
    # Alternatively, set the run's formatting explicitly
    # new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    # new_run.font.underline = True

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink

#This is only needed if you're using the builtin style above
def get_or_create_hyperlink_style(d):
    """If this document had no hyperlinks so far, the builtin
       Hyperlink style will likely be missing and we need to add it.
       There's no predefined value, different Word versions
       define it differently.
       This version is how Word 2019 defines it in the
       default theme, excluding a theme reference.
    """
    if "Hyperlink" not in d.styles:
        if "Default Character Font" not in d.styles:
            ds = d.styles.add_style("Default Character Font",
                                    docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                    True)
            ds.element.set(docx.oxml.shared.qn('w:default'), "1")
            ds.priority = 1
            ds.hidden = True
            ds.unhide_when_used = True
            del ds
        hs = d.styles.add_style("Hyperlink",
                                docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                True)
        hs.base_style = d.styles["Default Character Font"]
        hs.unhide_when_used = True
        hs.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
        hs.font.underline = True
        del hs

    return "Hyperlink"


def createworddoc(wordpath, readableposition, imagedirectory, hubheight, bladeradius):
    """
    Creates word document containing constraints images downloaded to 'imagedirectory'
    """

    document = Document()
    sections = document.sections
    for section in sections:
        margin = 1
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Open Sans Light'
    font.size = Pt(11)
    styles = document.styles
    style = styles.add_style('Attribution', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Open Sans Light'
    font.size = Pt(9)
    style = document.styles['Heading 1']
    font = style.font
    font.name = 'Open Sans ExtraBold'
    font.size = Pt(25)
    style = document.styles['Heading 2']
    font = style.font
    font.name = 'Open Sans ExtraBold'
    font.size = Pt(15)
    style = document.styles['Heading 3']
    font = style.font
    font.name = 'Open Sans Light'
    font.size = Pt(15)
    style = document.styles['List Bullet']
    font = style.font
    font.name = 'Open Sans ExtraBold'
    style = document.styles['List Bullet 2']
    style.left_margin = Cm(0)
    font = style.font
    font.name = 'Open Sans Light'

    p = document.add_paragraph(style='Heading 1')
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run('WeWantWind.org Turbine Siting Report')
    run.font.color.rgb = RGBColor.from_string('000000')
    p = document.add_paragraph(style='Heading 3')
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run('Position: ' + readableposition)
    p.paragraph_format.space_after = Pt(30)
    run.font.color.rgb = RGBColor.from_string('000000')

    image = imagedirectory + '/3d.png'
    document.add_picture(image, width=Cm(19))
    p = document.add_paragraph(style='Attribution')
    run = p.add_run('Satellite images © Esri — Source: Esri, i-cubed, USDA, USGS, AEX, GeoEye, Getmapping, Aerogrid, IGN, IGP, UPR-EGP, and the GIS User Community, ESRI')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    document.add_page_break()

    p = document.add_paragraph(style='Heading 1')
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run('WeWantWind.org Turbine Siting Report')
    run.font.color.rgb = RGBColor.from_string('000000')
    p = document.add_paragraph(style='Heading 3')
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run('Position: ' + readableposition)
    p.paragraph_format.space_after = Pt(30)
    run.font.color.rgb = RGBColor.from_string('000000')

    p = document.add_paragraph(style='Heading 2')
    run = p.add_run("Visibility map")
    run.font.color.rgb = RGBColor.from_string('000000')
    p = document.add_paragraph(style='List Bullet')
    run = p.add_run("Turbine hub viewable - hub height: " + str(hubheight) + "m")
    rgbarray = colors.to_rgb("#813FCB")
    run.font.color.rgb = RGBColor(int(255 * rgbarray[0]), int(255 * rgbarray[1]), int(255 * rgbarray[2]))
    p = document.add_paragraph(style='List Bullet')

    # p = document.add_paragraph(style='List Bullet 2')
    run = p.add_run("Turbine tip viewable - blade radius: " + str(bladeradius) + "m")
    rgbarray = colors.to_rgb("#B199CC")
    run.font.color.rgb = RGBColor(int(255 * rgbarray[0]), int(255 * rgbarray[1]), int(255 * rgbarray[2]))

    image = imagedirectory + '/viewshed.png'
    document.add_picture(image, width=Cm(19))
    p = document.add_paragraph(style='Attribution')
    run = p.add_run(" © ")
    add_hyperlink(p, 'OpenMapTiles', "https://www.openmaptiles.org")
    run = p.add_run(" © ")
    add_hyperlink(p, 'OpenStreetMap ', "https://www.openstreetmap.org/copyright")
    run = p.add_run("contributors")
    run.font.color.rgb = RGBColor.from_string('000000')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    document.add_page_break()

    p = document.add_paragraph('Below is a summary of all wind turbine planning constraints used to create your optimal wind turbine site.')

    lastcontraint = constraintslist[-1]
    for constraint in constraintslist:
        p = document.add_paragraph(style='Heading 2')
        run = p.add_run(constraint['heading'])
        run.font.color.rgb = RGBColor.from_string('000000')

        paragraph = document.add_paragraph('Constraints shown in this section:')
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        index = 0
        for dataset in constraint['datasets']:
            if constraint == constraintslist[0]: p = document.add_paragraph(style='List Bullet')
            else: p = document.add_paragraph(style='List Bullet 2')
            run = p.add_run(dataset)
            if constraint == constraintslist[0]:
                color = constraint['layers'][0]['color']
                if len(constraint['layers']) > 0: color = constraint['layers'][index]['color']
                rgbarray = colors.to_rgb(color)
                run.font.color.rgb = RGBColor(int(255 * rgbarray[0]), int(255 * rgbarray[1]), int(255 * rgbarray[2]))
            index += 1

        image = imagedirectory + '/' + constraint['heading'] + '.png'
        document.add_picture(image, width=Cm(19))
        p = document.add_paragraph(style='Attribution')
        run = p.add_run('Constraints data from multiple sources and copyright of respective data providers - for full list, go to ')
        add_hyperlink(p, 'ckan.wewantwind.org', "https://ckan.wewantwind.org")
        run.font.color.rgb = RGBColor.from_string('000000')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p = document.add_paragraph(style='Attribution')
        run = p.add_run(" © ")
        add_hyperlink(p, 'OpenMapTiles', "https://www.openmaptiles.org")
        run = p.add_run(" © ")
        add_hyperlink(p, 'OpenStreetMap ', "https://www.openstreetmap.org/copyright")
        run = p.add_run("contributors")
        run.font.color.rgb = RGBColor.from_string('000000')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        if constraint != lastcontraint:
            document.add_page_break()

    document.save(wordpath)

def createpdfdoc(pdfpath, readableposition, imagedirectory, hubheight, bladeradius):
    """
    Creates pdf document containing constraints images downloaded to 'imagedirectory'
    """

    canvas = Canvas(pdfpath, pagesize=A4)
    canvas.setTitle("WeWantWind.org Turbine Siting Report - " + readableposition)

    pdfmetrics.registerFont(TTFont('OpenSansLt', cwd + '/Open_Sans/static/OpenSans-Light.ttf'))
    pdfmetrics.registerFont(TTFont('OpenSans', cwd + '/Open_Sans/static/OpenSans-Medium.ttf'))
    pdfmetrics.registerFont(TTFont('OpenSansBd', cwd + '/Open_Sans/static/OpenSans-SemiBold.ttf'))
    pdfmetrics.registerFont(TTFont('OpenSansExtraBd', cwd + '/Open_Sans/static/OpenSans-ExtraBold.ttf'))

    # Cover page
    canvas.setFont("OpenSansExtraBd", 25) #choose your font type and font size
    canvas.drawString(40, 11*72, "WeWantWind.org Turbine Siting Report")
    canvas.setFont("OpenSansLt", 23) #choose your font type and font size
    canvas.drawString(40, 10.5*72, "Position: " + readableposition)
    image = imagedirectory + '/3d.png'
    canvas.drawInlineImage(image, 40 , 155, width=72*7,height=int(600 * 420 / 500))
    canvas.setFont("OpenSansLt", 8) #choose your font type and font size
    canvas.drawString(40, 142, 'Satellite images © Esri — Source: Esri, i-cubed, USDA, USGS, AEX, GeoEye, Getmapping, Aerogrid, IGN, IGP, UPR-EGP,')
    canvas.drawString(40, 130, 'and the GIS User Community, ESRI')
    canvas.showPage()

    # Viewshed
    canvas.setFont("OpenSansExtraBd", 25) #choose your font type and font size
    canvas.drawString(40, 11*72, "WeWantWind.org Turbine Siting Report")
    canvas.setFont("OpenSansLt", 23) #choose your font type and font size
    canvas.drawString(40, 10.5*72, "Position: " + readableposition)
    canvas.setFont("OpenSansExtraBd", 15) #choose your font type and font size
    canvas.drawString(40, 9.8*72, "Visibility map")
    vpos = 9.5*72
    canvas.setFont("OpenSansExtraBd", 11)
    rgbarray = colors.to_rgb("#813FCB")
    canvas.setFillColorRGB(rgbarray[0], rgbarray[1], rgbarray[2])    
    canvas.drawString(40, vpos, "• Turbine hub viewable - hub height: " + str(hubheight) + "m")
    vpos -= 16
    canvas.setFont("OpenSansExtraBd", 11)
    rgbarray = colors.to_rgb("#B199CC")
    canvas.setFillColorRGB(rgbarray[0], rgbarray[1], rgbarray[2])    
    canvas.drawString(40, vpos, "• Turbine tip viewable - blade radius: " + str(bladeradius) + "m")
    vpos -= 16
    image = imagedirectory + '/viewshed.png'
    canvas.drawInlineImage(image, 40 , 155, width=72*7,height=int(600 * 420 / 500))
    canvas.setFillColorRGB(0,0,0)
    canvas.setFont("OpenSansLt", 8) #choose your font type and font size
    canvas.drawString(40, vpos - (48 * 7) - 173, '© OpenMapTiles https://www.openmaptiles.org © OpenStreetMap contributors https://www.openstreetmap.org/copyright')
    canvas.showPage()

    # Constraints
    lastcontraint = constraintslist[-1]
    for constraint in constraintslist:
        canvas.setFont("OpenSansExtraBd", 25) #choose your font type and font size
        canvas.drawString(40, 11*72, "WeWantWind.org Turbine Siting Report")
        canvas.setFont("OpenSansLt", 23) #choose your font type and font size
        canvas.drawString(40, 10.5*72, "Position: " + readableposition)
        canvas.setFont("OpenSansExtraBd", 15) #choose your font type and font size
        canvas.drawString(40, 9.8*72, constraint['heading'])
        canvas.setFont("OpenSansLt", 11) #choose your font type and font size
        vpos = 9.5*72
        canvas.drawString(40, vpos, "Constraints shown in this section:")
        index = 0
        for dataset in constraint['datasets']:
            canvas.setFont("OpenSansLt", 11)
            canvas.setFillColorRGB(0,0,0)
            vpos -= 16
            if constraint == constraintslist[0]:
                canvas.setFont("OpenSansExtraBd", 11)
                color = constraint['layers'][0]['color']
                if len(constraint['layers']) > 0: color = constraint['layers'][index]['color']
                rgbarray = colors.to_rgb(color)
                canvas.setFillColorRGB(rgbarray[0], rgbarray[1], rgbarray[2])
            index += 1
            canvas.drawString(40, vpos, "• " + dataset)

        image = imagedirectory + '/' + constraint['heading'] + '.png'
        canvas.drawInlineImage(image, 40 , vpos - (48 * 7) - 100, width=72*7,height=420)
        canvas.setFont("OpenSansLt", 8)
        canvas.setFillColorRGB(0,0,0)
        canvas.drawString(40, vpos - (48 * 7) - 112, 'Constraints data from multiple sources and copyright of respective data providers - for full list, go to https://ckan.wewantwind.org')
        canvas.drawString(40, vpos - (48 * 7) - 124, '© OpenMapTiles https://www.openmaptiles.org © OpenStreetMap contributors https://www.openstreetmap.org/copyright')

        if constraint != lastcontraint:
            canvas.showPage()

    canvas.save()

def GetReport(type, lat, lng, hubheight, bladeradius):
    """
    Return report object for specific type, lat, lng
    """

    id = str(uuid.uuid4())
    lat, lng = round(lat, COORDINATE_PRECISION), round(lng, COORDINATE_PRECISION)
    filestem = str(lat) + "_" + str(lng) + "_" + str(hubheight) + "_" + str(bladeradius)
    readableposition = str(lat) + "°N, " + str(lng) + "°E"
    downloadsdirectory = cwd + '/downloads/'
    pdfpath = downloadsdirectory + filestem + '.pdf'
    wordpath = downloadsdirectory + filestem + '.docx'
    # if True:
    if (os.path.isfile(wordpath) is False) or (os.path.isfile(pdfpath) is False):
        imagedirectory = processimages(id, [lng, lat], constraintslist, {'width': '600', 'height': '500', 'ratio': '3', 'zoom': '12', 'center': str(lng) + ',' + str(lat)}, hubheight, bladeradius)
        if os.path.isfile(wordpath) is False: createworddoc(wordpath, readableposition, imagedirectory, hubheight, bladeradius)
        if os.path.isfile(pdfpath) is False: createpdfdoc(pdfpath, readableposition, imagedirectory, hubheight, bladeradius)
        shutil.rmtree(imagedirectory)

    returnfile = wordpath
    filestem = "WeWantWind.org Report - " + readableposition
    mimetype = 'application/msword'
    filename = filestem + ".docx"
    if type == 'pdf': 
        returnfile = pdfpath
        mimetype = 'application/pdf'
        filename = filestem + ".pdf"

    try:    
        with open(returnfile, 'rb') as f:
           file_data = f.read()
        response = HttpResponse(file_data, content_type=mimetype)
        response['Content-Disposition'] = 'attachment; filename="' + filename + '"'

    except IOError:
        # handle file not exist case here
        response = HttpResponseNotFound('<h1>File does not exist</h1>')

    return response

@csrf_exempt
def SiteReport(request):
    """
    Return site report as docs or pdf
    """

    type = request.GET.get('type','word')
    lat = float(request.GET.get('lat', 51))
    lng = float(request.GET.get('lng',0))
    hubheight = float(request.GET.get('hub', DEFAULT_HUB_HEIGHT))
    bladeradius = float(request.GET.get('blade', DEFAULT_BLADE_RADIUS))

    return GetReport(type, lat, lng, hubheight, bladeradius)

@csrf_exempt
def NearestTurbineReport(request):
    type = request.GET.get('type','word')
    lat = float(request.GET.get('lat', 51))
    lng = float(request.GET.get('lng',0))
    hubheight = float(request.GET.get('hub', DEFAULT_HUB_HEIGHT))
    bladeradius = float(request.GET.get('blade', DEFAULT_BLADE_RADIUS))

    site = GetNearestTurbine(lat, lng)
 
    return GetReport(type, site['turbinelat'], site['turbinelng'], hubheight, bladeradius)

@csrf_exempt
def CreateGeoJSON(request):
    lat = float(request.GET.get('lat', 0))
    lng = float(request.GET.get('lng',0))
    geojson = {
        "type": "FeatureCollection",
        "features": [
            {
            "type": "Feature",
            "name": "WeWantWind.org Turbine",
            "geometry": {
                "type": "Point",
                "coordinates": [lng, lat]
            },
        }]
    }

    return OutputJson(geojson)

@csrf_exempt
def CastVote(request):
    """
    Cast provisional vote and sent email requesting vote confirmation
    """

    try:
        parameters = json.loads(request.body)
    except ValueError:
        return OutputError()

    ''' Begin reCAPTCHA validation '''

    recaptcha_response = parameters['recaptcha']
    url = 'https://www.google.com/recaptcha/api/siteverify'
    values = {
        'secret': GOOGLE_RECAPTCHA_SECRET_KEY,
        'response': recaptcha_response
    }
    recaptchadata = urllib.parse.urlencode(values).encode()
    req =  urllib.request.Request(url, data=recaptchadata)
    response = urllib.request.urlopen(req)
    result = json.loads(response.read().decode())

    ''' End reCAPTCHA validation '''
    ip, is_routable = get_client_ip(request)

    if result['success']:
        token = uuid.uuid4().hex
        # Save vote in database
        provisionalvote = Vote( \
            name=parameters['name'], \
            email=parameters['email'].lower(), \
            contactable=parameters['contactable'], \
            userlocation=Point(parameters['userlocation']['lng'], parameters['userlocation']['lat']), \
            site=Point(parameters['site']['lng'], parameters['site']['lat']), \
            ip=ip, \
            token=token)
        provisionalvote.save()
        # Attempt to send email
        from_email = '"WeWantWind.org" <info@wewantwind.org>'
        subject = "WeWantWind.org: Confirm your wind turbine vote"
        current_site = get_current_site(request)
        parameters['domain'] = current_site.domain
        parameters['uid'] = urlsafe_base64_encode(force_bytes(provisionalvote.pk))
        parameters['token'] = token
        parameters['site'] = {'lat': round(parameters['site']['lat'], COORDINATE_PRECISION), 'lng': round(parameters['site']['lng'], COORDINATE_PRECISION)}
        message = render_to_string('backend/confirm_vote.html', parameters)        
        message = EmailMessage(subject, message, from_email=from_email, to=[parameters['email']])
        message.send()
    return OutputJson({'result': 'success'})

@csrf_exempt
def ConfirmVote(request, uidb64, token):
    """
    Confirm vote using link sent via email
    """

    sleep(0.5)

    try:
        id = urlsafe_base64_decode(uidb64).decode()
        provisionalvote = Vote.objects.get(pk=id)
    except (TypeError, ValueError, OverflowError, Vote.DoesNotExist):
        provisionalvote = None

    if (provisionalvote is not None) and (provisionalvote.token == token) and (provisionalvote.confirmed == False):
        Vote.objects.filter(email=provisionalvote.email).filter(~Q(pk=id)).delete()
        provisionalvote.confirmed = True
        provisionalvote.save()
        return render(request, 'backend/vote_confirmed.html')
    else:
        return render(request, 'backend/vote_not_confirmed.html')

@csrf_exempt
def Votes(request):
    """
    Get data on all votes
    """

    distinctpoints = Vote.objects.filter(confirmed=True).values('site').annotate(Count('id')).order_by()

    features = []
    for distinctpoint in distinctpoints:
        # print(distinctpoint, distinctpoint['site'].coords)
        allvotes = Vote.objects.filter(confirmed=True).filter(site=distinctpoint['site']).count()
        onemilevotes = Vote.objects.filter(confirmed=True).filter(site=distinctpoint['site']).filter(userlocation__distance_lte=(distinctpoint['site'], RadiusDistance(mi=1))).count()
        fivemilevotes = Vote.objects.filter(confirmed=True).filter(site=distinctpoint['site']).filter(userlocation__distance_lte=(distinctpoint['site'], RadiusDistance(mi=5))).count()
        tenmilevotes = Vote.objects.filter(confirmed=True).filter(site=distinctpoint['site']).filter(userlocation__distance_lte=(distinctpoint['site'], RadiusDistance(mi=10))).count()
        feature = {
            "type": "feature",
            "properties": {
                'name': 'Site votes', 
                'subtype': 'votes',
                'position': str(round(distinctpoint['site'].coords[1], COORDINATE_PRECISION)) + "°N, " + str(round(distinctpoint['site'].coords[0], COORDINATE_PRECISION)) + "°E",
                'votes': allvotes,
                'lat': distinctpoint['site'].coords[1],
                'lng': distinctpoint['site'].coords[0],
                'votes:within:1:mile': onemilevotes,
                'votes:within:5:miles': fivemilevotes,
                'votes:within:10:miles': tenmilevotes
            },
            "geometry": {
                "type": "Point",
                "coordinates": [distinctpoint['site'].coords[0], distinctpoint['site'].coords[1]]
            }
        }
        features.append(feature)

    geojson = { "type": "FeatureCollection", "features": features }
    return OutputJson(geojson)

@csrf_exempt
def LocalPeople(request):
    """
    Get number of local people within a predefined distance of supplied position
    """

    try:
        parameters = json.loads(request.body)
    except ValueError:
        return OutputError()
    
    centre = Point(parameters['lng'], parameters['lat'], srid=4326)    
    localpeople = Vote.objects.filter(confirmed=True, contactable=True, userlocation__distance_lte=(centre, RadiusDistance(mi=LOCAL_DISTANCE)))
    localpeoplecount = 0
    if localpeople is not None: localpeoplecount = len(localpeople)
    return OutputJson({'localpeople': localpeoplecount})

@csrf_exempt
def SendMessage(request):
    """
    Create provisional message entry
    """

    try:
        parameters = json.loads(request.body)
    except ValueError:
        return OutputError()

    ''' Begin reCAPTCHA validation '''

    recaptcha_response = parameters['recaptcha']
    url = 'https://www.google.com/recaptcha/api/siteverify'
    values = {
        'secret': GOOGLE_RECAPTCHA_SECRET_KEY,
        'response': recaptcha_response
    }
    recaptchadata = urllib.parse.urlencode(values).encode()
    req =  urllib.request.Request(url, data=recaptchadata)
    response = urllib.request.urlopen(req)
    result = json.loads(response.read().decode())

    ''' End reCAPTCHA validation '''
    ip, is_routable = get_client_ip(request)

    if result['success']:
        token = uuid.uuid4().hex
        # Save vote in database
        provisionalmessage = Message( \
            name=parameters['name'], \
            email=parameters['email'].lower(), \
            userlocation=Point(parameters['userlocation']['lng'], parameters['userlocation']['lat']), \
            ip=ip, \
            token=token)
        provisionalmessage.save()
        # Attempt to send email
        from_email = '"WeWantWind.org" <info@wewantwind.org>'
        subject = "WeWantWind.org: Confirm your message request"
        current_site = get_current_site(request)
        parameters['domain'] = current_site.domain
        parameters['uid'] = urlsafe_base64_encode(force_bytes(provisionalmessage.pk))
        parameters['token'] = token
        parameters['localdistance'] = LOCAL_DISTANCE
        message = render_to_string('backend/confirm_message.html', parameters)        
        message = EmailMessage(subject, message, from_email=from_email, to=[parameters['email']])
        message.send()
    return OutputJson({'result': 'success'})

@csrf_exempt
def SendShare(request):
    """
    Send shareable link to email
    """

    try:
        parameters = json.loads(request.body)
    except ValueError:
        return OutputError()

    ''' Begin reCAPTCHA validation '''

    recaptcha_response = parameters['recaptcha']
    url = 'https://www.google.com/recaptcha/api/siteverify'
    values = {
        'secret': GOOGLE_RECAPTCHA_SECRET_KEY,
        'response': recaptcha_response
    }
    recaptchadata = urllib.parse.urlencode(values).encode()
    req =  urllib.request.Request(url, data=recaptchadata)
    response = urllib.request.urlopen(req)
    result = json.loads(response.read().decode())

    ''' End reCAPTCHA validation '''
    ip, is_routable = get_client_ip(request)

    if result['success']:
        # Attempt to send email
        from_email = '"WeWantWind.org" <info@wewantwind.org>'
        subject = "WeWantWind.org: Someone has shared a wind turbine site link with you"
        current_site = get_current_site(request)
        parameters['domain'] = current_site.domain
        message = render_to_string('backend/share.html', parameters)        
        message = EmailMessage(subject, message, from_email=from_email, to=[parameters['email']])
        message.send()
    return OutputJson({'result': 'success'})

@csrf_exempt
def ProcessMessageQueue(request):
    """
    Process message queue
    Run this once a day - around 8am - to prevent users getting deluged with emails
    """

    current_site = get_current_site(request)
    from_email = '"WeWantWind.org" <info@wewantwind.org>'
    subject = "WeWantWind.org: Introductory email from user(s) wanting to connect with other users"
    pendingmessages = Message.objects.filter(sent=False)
    outboundqueue = {}
    for pendingmessage in pendingmessages:
        print("Sending message from", pendingmessage.name, pendingmessage.ip)
        parameters = {\
            'domain': current_site.domain, \
            'sourcename': pendingmessage.name, \
            'sourceemail': pendingmessage.email \
        }
        localpeople = Vote.objects.filter(~Q(email=pendingmessage.email)).filter(confirmed=True, contactable=True, userlocation__distance_lte=(pendingmessage.userlocation, RadiusDistance(mi=LOCAL_DISTANCE)))
        for localperson in localpeople:
            name, email = localperson.name, localperson.email
            if email not in outboundqueue:
                outboundqueue[email] = {'name': name, 'pk': localperson.pk, 'people': []}
            outboundqueue[email]['people'].append({'name': pendingmessage.name, 'email': pendingmessage.email})

    messagessent = 0
    for email in outboundqueue:
        name = outboundqueue[email]['name']
        peoplelist = ''
        for person in outboundqueue[email]['people']:
            peoplelist += person['name'].strip() + ": " + person['email'].strip() + "\n"
        if peoplelist != '':
            parameters = {'domain': current_site.domain, 'name': name, 'peoplelist': peoplelist}
            parameters['uid'] = urlsafe_base64_encode(force_bytes(outboundqueue[email]['pk']))
            parameters['token'] = uuid.uuid4().hex
            parameters['localdistance'] = LOCAL_DISTANCE
            Vote.objects.filter(pk=outboundqueue[email]['pk']).update(token=parameters['token'])
            message = render_to_string('backend/standard_message.html', parameters)        
            message = EmailMessage(subject, message, from_email=from_email, to=[email])
            message.send()
            messagessent += 1

    # pprint(outboundqueue, indent=4)            
    Message.objects.filter(sent=False).update(sent=True)

    return OutputJson({'result': 'success', 'messagesent': messagessent})

@csrf_exempt
def ConfirmMessage(request, uidb64, token):
    """
    Confirm vote using link sent via email
    """

    sleep(0.5)

    try:
        id = urlsafe_base64_decode(uidb64).decode()
        provisionalmessage = Message.objects.get(pk=id)
    except (TypeError, ValueError, OverflowError, Message.DoesNotExist):
        provisionalmessage = None

    if (provisionalmessage is not None) and (provisionalmessage.token == token) and (provisionalmessage.confirmed == False):
        Message.objects.filter(email=provisionalmessage.email).filter(~Q(pk=id)).delete()
        provisionalmessage.confirmed = True
        provisionalmessage.save()
        return render(request, 'backend/message_confirmed.html')
    else:
        return render(request, 'backend/message_not_confirmed.html')

@csrf_exempt
def RemoveMailingList(request, uidb64, token):
    """
    Remove from mailing list
    """

    sleep(0.5)

    try:
        pk = urlsafe_base64_decode(uidb64).decode()
        provisionalvote = Vote.objects.get(pk=pk)
    except (TypeError, ValueError, OverflowError, Vote.DoesNotExist):
        provisionalvote = None

    if (provisionalvote is not None) and (provisionalvote.token == token) and (provisionalvote.contactable == True):
        Message.objects.filter(email=provisionalvote.email).filter(~Q(pk=pk)).delete()
        provisionalvote.contactable = False
        provisionalvote.save()
        return render(request, 'backend/mailinglist_confirmed.html')
    else:
        return render(request, 'backend/mailinglist_not_confirmed.html')

@csrf_exempt
def Test(request):
    return render(request, 'backend/vote_not_confirmed.html')

class HttpResponseTemporaryRedirect(HttpResponse):
    status_code = 307

    def __init__(self, redirect_to):
        HttpResponse.__init__(self)
        self['Location'] = iri_to_uri(redirect_to)

class IntentSchemeRedirect(HttpResponseTemporaryRedirect):
    allowed_schemes = ['intent']

@csrf_exempt
def Shortcode(request, code):
    return IntentSchemeRedirect("intent://#Intent;S.data=https://wewantwind.org/static/geojson/dummyturbine3.geojson;scheme=windview;package=org.wewantwind.windview;end")

def returncirclesforpoint(lng, lat):
    center = GeoJSONFeature(geometry=GeoJSONPoint((lng, lat)))
    bearing1 = 0
    bearing2 = 359.99999

    features = []
    steps = [5, 10, 15, 20, 25, 30, 35, 40]
    stepdistance = 10
    for step_index in range(len(steps)):
        radius = steps[step_index]
        feature = line_arc(center=center, radius=radius, bearing1=bearing1, bearing2=bearing2)
        feature['properties'] = {'class': 'Distance_Circle', 'distance': str(radius) + 'km'}
        features.append(feature)
        feature_coordinates = feature['geometry']['coordinates']
        point_label_coordinates = feature_coordinates[int(len(feature_coordinates) / 2)]
        feature_point = {'type': 'Feature', 'name': str(radius) + 'km', 'properties': {'name': str(radius) + 'km', 'class': 'Distance_Circle_Label'}, 'geometry': {'type': 'Point', 'coordinates': point_label_coordinates}}
        features.append(feature_point)

    return features

def getelevationforpoint(lon, lat):
    global TERRAIN_FILE
    # With thanks to https://stackoverflow.com/questions/74026802/get-elevation-from-lat-long-of-geotiff-data-in-gdal
    ds = gdal.OpenEx(TERRAIN_FILE)
    raster_proj = ds.GetProjection()
    gt = ds.GetGeoTransform()
    ds = None
    source_srs = osr.SpatialReference()
    source_srs.ImportFromWkt(osr.GetUserInputAsWKT("urn:ogc:def:crs:OGC:1.3:CRS84"))
    target_srs = osr.SpatialReference()
    target_srs.ImportFromWkt(raster_proj)
    ct = osr.CoordinateTransformation(source_srs, target_srs)
    mapx, mapy, *_ = ct.TransformPoint(lon, lat)
    gt_inv = gdal.InvGeoTransform(gt) 
    px, py = gdal.ApplyGeoTransform(gt_inv, mapx, mapy)
    py = int(py)
    px = int(px)
    ds = gdal.OpenEx(TERRAIN_FILE)
    elevation_value = ds.ReadAsArray(px, py, 1, 1)
    ds = None
    elevation = elevation_value[0][0]
    return elevation, mapx, mapy

def cropGeoTIFF(lon, lat, distance, inputfile, outputfile):
    ds = gdal.OpenEx(inputfile)
    raster_proj = ds.GetProjection()
    gt = ds.GetGeoTransform()
    ds = None
    source_srs = osr.SpatialReference()
    source_srs.ImportFromWkt(osr.GetUserInputAsWKT("urn:ogc:def:crs:OGC:1.3:CRS84"))
    target_srs = osr.SpatialReference()
    target_srs.ImportFromWkt(raster_proj)
    ct = osr.CoordinateTransformation(source_srs, target_srs)
    mapx, mapy, *_ = ct.TransformPoint(lon, lat)

    upper_left_x = mapx - distance
    upper_left_y = mapy + distance
    lower_right_x = mapx + distance
    lower_right_y = mapy - distance
    window = (upper_left_x,upper_left_y,lower_right_x,lower_right_y)

    gdal.Translate(outputfile, inputfile, projWin = window)

def subtractimage(image_main, image_to_subtract):
    image = Image.open(image_main)
    subtractimage = Image.open(image_to_subtract)
    buffer1    = np.asarray(image)
    buffer2    = np.asarray(subtractimage)
    buffer3    = buffer1 - buffer2
    image     = Image.fromarray(buffer3)
    image.save(image_main)

def createcolourizedtransparentPNG(inputfile, outputfile, new_color, clippingimage=None):
    image = Image.open(inputfile)
    if clippingimage is not None:
        subtractimage = Image.open(clippingimage)
        buffer1    = np.asarray(image)
        buffer2    = np.asarray(subtractimage)
        buffer3    = buffer1 - buffer2
        image     = Image.fromarray(buffer3)
    mask = image.copy()
    mask = mask.convert("L")
    pixels = list(image.getdata())
    modified_pixels = [new_color if pixel != (255, 255, 255) else pixel for pixel in pixels]
    modified_image = Image.new("RGBA", image.size)
    modified_image.putdata(modified_pixels)
    modified_image.putalpha(mask)
    modified_image.save(outputfile)

def read_file(filename):
    vsifile = gdal.VSIFOpenL(filename,'r')
    gdal.VSIFSeekL(vsifile, 0, 2)
    vsileng = gdal.VSIFTellL(vsifile)
    gdal.VSIFSeekL(vsifile, 0, 0)
    return gdal.VSIFReadL(1, vsileng, vsifile)

def GetViewsheds(lon, lat, hubheight, bladeradius):
    global TERRAIN_FILE

    uniqueid = str(lon) + '_' + str(lat) + '_' + str(hubheight) + '_' + str(bladeradius)
    groundheight, observerX, observerY = getelevationforpoint(lon, lat)
    towerheight = groundheight + hubheight
    turbinetip = groundheight + hubheight + bladeradius
    towerheight_outfile = uniqueid + "_tower.tif"
    turbinetip_outfile = uniqueid + "_tip.tif"
    towerheight_outfile = '/vsimem/' + uniqueid + "_tower.tif"
    turbinetip_outfile = '/vsimem/' + uniqueid + "_tip.tif"
    # towerheight_png = uniqueid + "_tower.png"
    # turbinetip_png = uniqueid + "_tip.png"
    viewshedcolor_towerheight = (165, 147, 171, 50)  
    viewshedcolor_turbinetip = (211, 207, 225, 50)  

    # ds = gdal.ViewshedGenerate(
    #     srcBand = src_ds.GetRasterBand(1),
    #     driverName = 'GTiff',
    #     targetRasterName = "test.tif",
    #     creationOptions = ["INTERLEAVE=BAND"],
    #     observerX = observerX,
    #     observerY = observerY,
    #     observerHeight = 1.5,
    #     targetHeight = 0,  
    #     visibleVal = 255,  
    #     invisibleVal = 0,  
    #     outOfRangeVal = 0,  
    #     noDataVal = -1.0,  
    #     dfCurvCoeff = 0.85714,
    #     mode = gdal.GVM_Edge,
    #     maxDistance = VIEWSHED_MAX_DISTANCE,
    #     heightMode = gdal.GVOT_MIN_TARGET_HEIGHT_FROM_GROUND,
    #     options = ["UNUSED=YES"],
    # )

    src_ds = gdal.Open(TERRAIN_FILE)

    gdal.ViewshedGenerate(
        srcBand = src_ds.GetRasterBand(1),
        driverName = 'GTiff',
        targetRasterName = towerheight_outfile,
        creationOptions = [],
        observerX = observerX,
        observerY = observerY,
        observerHeight = int(towerheight + 0.5),
        targetHeight = 1.5,
        visibleVal = 255.0,
        invisibleVal = 0.0,
        outOfRangeVal = 0.0,
        noDataVal = 0.0,
        dfCurvCoeff = 1.0,
        mode = 1,
        maxDistance = VIEWSHED_MAX_DISTANCE) 

    gdal.ViewshedGenerate(
        srcBand = src_ds.GetRasterBand(1),
        driverName = 'GTiff',
        targetRasterName = turbinetip_outfile,
        creationOptions = [],
        observerX = observerX,
        observerY = observerY,
        observerHeight = int(turbinetip + 0.5),
        targetHeight = 1.5,
        visibleVal = 255.0,
        invisibleVal = 0.0,
        outOfRangeVal = 0.0,
        noDataVal = 0.0,
        dfCurvCoeff = 1.0,
        mode = 1,
        maxDistance = VIEWSHED_MAX_DISTANCE) 

    all_features = distancecircles = returncirclesforpoint(lon, lat)    
    towerheight_geojson = json.loads(polygonizeraster(uniqueid, towerheight_outfile))
    turbinetip_geojson = json.loads(polygonizeraster(uniqueid, turbinetip_outfile))

    for feature_index in range(len(towerheight_geojson['features'])):
        towerheight_geojson['features'][feature_index]['properties']['class'] = 'viewshed_towerheight'
        all_features.append(towerheight_geojson['features'][feature_index])
    for feature_index in range(len(turbinetip_geojson['features'])):
        turbinetip_geojson['features'][feature_index]['properties']['class'] = 'viewshed_turbinetip'
        all_features.append(turbinetip_geojson['features'][feature_index])
    
    featurecollection = {'type': 'FeatureCollection', 'features': all_features}

    return featurecollection

def reprojectrasterto4326(input_file, output_file):
    warp = gdal.Warp(output_file, gdal.Open(input_file), dstSRS='EPSG:4326')
    warp = None

def polygonizeraster(uniqueid, raster_file):
    memory_geojson = '/vsimem/' + uniqueid + ".geojson"
    memory_transformed_raster = '/vsimem/' + uniqueid + '.tif'
    reprojectrasterto4326(raster_file, memory_transformed_raster)

    driver = ogr.GetDriverByName("GeoJSON")
    ds = gdal.OpenEx(memory_transformed_raster)
    raster_proj = ds.GetProjection()
    ds = None
    source_srs = osr.SpatialReference()
    source_srs.ImportFromWkt(raster_proj)
    src_ds = gdal.Open(memory_transformed_raster)
    srs = osr.SpatialReference()
    srs.ImportFromWkt(src_ds.GetProjection())    
    srcband = src_ds.GetRasterBand(1)

    dst_ds = driver.CreateDataSource(memory_geojson)
    dst_layer = dst_ds.CreateLayer("viewshed", srs = source_srs)
    newField = ogr.FieldDefn('Area', ogr.OFTInteger)
    dst_layer.CreateField(newField)
    polygonize = gdal.Polygonize(srcband, srcband, dst_layer, 0, [], callback=None )
    polygonize = None
    del dst_ds

    geojson_content = read_file(memory_geojson)

    return geojson_content

@csrf_exempt
def Viewshed(request):
    """
    Return viewshed as GeoJSON
    """
    parameters, lat, lng = None, None, None

    try:
        parameters = json.loads(request.body)
        lat = float(parameters.get('lat', 51))
        lng = float(parameters.get('lng',0))
        hubheight = float(parameters.get('hub', DEFAULT_HUB_HEIGHT))
        bladeradius = float(parameters.get('blade', DEFAULT_BLADE_RADIUS))
    except ValueError:
        lat = request.GET.get('lat', None)
        lng = request.GET.get('lng', None)        
        if (lat is None) or (lng is None):
            return OutputError()
        lat = float(lat)
        lng = float(lng)
        hubheight = float(request.GET.get('hub', DEFAULT_HUB_HEIGHT))
        bladeradius = float(request.GET.get('blade', DEFAULT_BLADE_RADIUS))        

    geojson = GetViewsheds(lng, lat, hubheight, bladeradius)

    return OutputJson(geojson)

