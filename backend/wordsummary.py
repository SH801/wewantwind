import os
from docx import Document
import docx
import json
import requests
import shutil
import uuid
from matplotlib import colors
from pprint import pprint
from docx.shared import Cm, Pt, RGBColor
from reportlab.pdfgen.canvas import Canvas
from reportlab.lib.units import cm, inch
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

cwd = os.path.dirname(os.path.realpath(__file__))

constraintslist = [
    {
        'heading': 'All constraints',
        'datasets': [
            'Inadequate wind speed',
            'Landscape and visual impact',
            'Heritage impacts',
            'Separation distance to residential properties',
            'Ecology and wildlife',
            'Other technical constraints',
            'Aviation and exclusion areas'
        ],
        'layers': [
            {'name': 'wind', 'color': 'blue'},
            {'name': 'landscape_and_visual_impact', 'color': 'chartreuse'},
            {'name': 'heritage_impacts', 'color': 'darkgoldenrod'},
            {'name': 'separation_distance_to_residential', 'color': 'darkorange'},
            {'name': 'ecology_and_wildlife', 'color': 'darkgreen'},
            {'name': 'roads', 'color': 'red'},
            {'name': 'aviation_and_exclusion_areas', 'color': 'purple'}
        ]
    },
    {
        'heading': 'Inadequate wind speed', 
        'datasets': [
            'Inadequate wind speed (< 5 metres per second) from NOABL wind speed database'
        ],
        'layers': [
            {'name': 'wind', 'color': 'blue'}
        ]
    },
    {
        'heading': 'Landscape and visual impact', 
        'datasets': [
            'National Parks', 
            'Areas of Outstanding Natural Beauty / National Scenic Areas', 
            'Heritage Coasts'
        ],
        'layers': [
            {'name': 'landscape_and_visual_impact', 'color': 'chartreuse'}
        ]
    },
    {
        'heading': 'Heritage impacts', 
        'datasets': [
            'Listed buildings',
            'Conservation areas',
            'World Heritage Sites',
            'Scheduled Ancient Monuments',
            'Registered parks and gardens',
            'Registered historic battlefields'
        ],
        'layers': [
            {'name': 'heritage_impacts', 'color': 'darkgoldenrod'}
        ]
    },
    {
        'heading': 'Separation distance to residential properties', 
        'datasets': [
            '400 metre buffer from residential areas'
        ],
        'layers': [
            {'name': 'separation_distance_to_residential', 'color': 'darkorange'}
        ]
    },
    {
        'heading': 'Ecology and wildlife', 
        'datasets': [
            'Sites of Special Scientific Interest / Areas of Special Scientific Interest',
            'Ramsar sites',
            'Special Areas of Conservation',
            'Special Protection Areas',
            'National nature reserves',
            'Ancient woodland',
            'Local wildlife reserves',
            '50 metre buffer from hedgerows'
        ],
        'layers': [
            {'name': 'ecology_and_wildlife', 'color': 'darkgreen'}
        ]
    },
    {
        'heading': 'Other technical constraints', 
        'datasets': [
            '150 metre buffer from public roads (A and B roads and motorways)',
            '150 metre buffer from railway lines',
            '150 metre buffer from inland waters',
            '150 metre buffer from pipelines',
            '150 metre buffer from power lines',
            '150 metre buffer from public footpaths',
            '200 metre buffer from bridleways',     
        ],
        'layers': [
            {'name': 'roads', 'color': 'red'}
        ]
    },
    {
        'heading': 'Aviation and exclusion areas', 
        'datasets': [
            'Civilian airports',
            'Aerodromes',
            'Military airfields and airbases',
            'MOD training areas',
            'Explosive safeguarded areas, danger areas near ranges',
            'MOD exclusion areas'
        ],
        'layers': [
            {'name': 'aviation_and_exclusion_areas', 'color': 'purple'}
        ]
    }
]

def processimages(id, coordinates, constraintslist, parameters):
    with open(cwd + '/styles/planningconstraint.json') as fp: planningconstraint = json.load(fp)
    with open(cwd + '/styles/planningconstraints.json') as fp: planningconstraints = json.load(fp)

    imagedirectory = cwd + '/downloads/' + id
    os.mkdir(imagedirectory) 

    # Modify planningconstraint template according to constraint color and duplicating for each listed layer
    for constraint in constraintslist:
        print(constraint['heading'])
        alllayers = []
        constraintlayers = json.loads(json.dumps(planningconstraint))
        for layer in constraint['layers']:
            for constraintlayer in constraintlayers:
                newconstraintlayer = json.loads(json.dumps(constraintlayer))
                newconstraintlayer['id'] = layer['name'] + "_" + newconstraintlayer['id']
                newconstraintlayer['source-layer'] = layer['name']
                if 'fill-color' in newconstraintlayer['paint']: newconstraintlayer['paint']['fill-color'] = layer['color']
                alllayers.append(newconstraintlayer)

        # Insert new layers into existing stylesheet template
        newlayers = []
        newplanningconstraints = json.loads(json.dumps(planningconstraints))
        for newplanningconstraintlayer in newplanningconstraints['style']['layers']:
            if newplanningconstraintlayer['id'] == 'planningconstraints':
                for layer in alllayers:
                    newlayers.append(layer)
            else: 
                newlayers.append(newplanningconstraintlayer)

        if 'customgeojson' in newplanningconstraints['style']['sources']:
            newplanningconstraints['style']['sources']['customgeojson'] = {
                "type": "geojson",
                "data": {"type": "FeatureCollection", "features": [{
                    "type": "Feature",
                    "geometry": {
                        "type": "Point",
                        "coordinates": coordinates
                    }
                }]}
            }      

        newplanningconstraints['style']['layers'] = newlayers
        for key in parameters: newplanningconstraints[key] = parameters[key]

        # pprint(newplanningconstraints, indent=4)
        r = requests.post('http://localhost:81/render', json=newplanningconstraints, stream=True)
        if r.status_code == 200:
            with open(imagedirectory + "/" + constraint['heading'] + '.png', 'wb') as f:
                r.raw.decode_content = True
                shutil.copyfileobj(r.raw, f) 

    return imagedirectory

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(
        docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    new_run.style = get_or_create_hyperlink_style(part.document)
    # Alternatively, set the run's formatting explicitly
    # new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    # new_run.font.underline = True

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink

#This is only needed if you're using the builtin style above
def get_or_create_hyperlink_style(d):
    """If this document had no hyperlinks so far, the builtin
       Hyperlink style will likely be missing and we need to add it.
       There's no predefined value, different Word versions
       define it differently.
       This version is how Word 2019 defines it in the
       default theme, excluding a theme reference.
    """
    if "Hyperlink" not in d.styles:
        if "Default Character Font" not in d.styles:
            ds = d.styles.add_style("Default Character Font",
                                    docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                    True)
            ds.element.set(docx.oxml.shared.qn('w:default'), "1")
            ds.priority = 1
            ds.hidden = True
            ds.unhide_when_used = True
            del ds
        hs = d.styles.add_style("Hyperlink",
                                docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                True)
        hs.base_style = d.styles["Default Character Font"]
        hs.unhide_when_used = True
        hs.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
        hs.font.underline = True
        del hs

    return "Hyperlink"


def createworddoc(wordpath, imagedirectory):
    document = Document()
    sections = document.sections
    for section in sections:
        margin = 1
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Open Sans Light'
    font.size = Pt(9)
    style = document.styles['Heading 1']
    font = style.font
    font.name = 'Open Sans ExtraBold'
    font.size = Pt(25)
    style = document.styles['Heading 2']
    font = style.font
    font.name = 'Open Sans Extra Bold'
    font.size = Pt(15)
    style = document.styles['List Bullet']
    font = style.font
    font.name = 'Open Sans ExtraBold'
    style = document.styles['List Bullet 2']
    style.left_margin = Cm(0)
    font = style.font
    font.name = 'Open Sans Light'

    p = document.add_paragraph(style='Heading 1')
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run('Wewantwind Turbine Siting Report')
    run.font.color.rgb = RGBColor.from_string('000000')
    p = document.add_paragraph('Below is a summary of all wind turbine planning constraints used to create your optimal wind turbine site.')

    lastcontraint = constraintslist[-1]
    for constraint in constraintslist:
        p = document.add_paragraph(style='Heading 2')
        run = p.add_run(constraint['heading'])
        run.font.color.rgb = RGBColor.from_string('000000')

        paragraph = document.add_paragraph('Constraints shown in this section:')
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        index = 0
        for dataset in constraint['datasets']:
            if constraint == constraintslist[0]: p = document.add_paragraph(style='List Bullet')
            else: p = document.add_paragraph(style='List Bullet 2')
            run = p.add_run(dataset)
            if constraint == constraintslist[0]:
                color = constraint['layers'][0]['color']
                if len(constraint['layers']) > 0: color = constraint['layers'][index]['color']
                rgbarray = colors.to_rgb(color)
                run.font.color.rgb = RGBColor(int(255 * rgbarray[0]), int(255 * rgbarray[1]), int(255 * rgbarray[2]))
            index += 1

        image = imagedirectory + '/' + constraint['heading'] + '.png'
        document.add_picture(image, width=Cm(19))
        paragraph = document.add_paragraph('Constraints data from multiple sources and copyright of respective data providers - for full list refer to ')
        add_hyperlink(paragraph, 'ckan.wewantwind.org', "https://ckan.wewantwind.org")
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

        if constraint != lastcontraint:
            document.add_page_break()

    document.save(wordpath)

def createpdfdoc(pdfpath, imagedirectory):
    canvas = Canvas(pdfpath, pagesize=A4)

    pdfmetrics.registerFont(TTFont('OpenSansLt', cwd + '/Open_Sans/static/OpenSans-Light.ttf'))
    pdfmetrics.registerFont(TTFont('OpenSans', cwd + '/Open_Sans/static/OpenSans-Medium.ttf'))
    pdfmetrics.registerFont(TTFont('OpenSansBd', cwd + '/Open_Sans/static/OpenSans-SemiBold.ttf'))
    pdfmetrics.registerFont(TTFont('OpenSansExtraBd', cwd + '/Open_Sans/static/OpenSans-ExtraBold.ttf'))

    lastcontraint = constraintslist[-1]
    for constraint in constraintslist:
        canvas.setFont("OpenSansExtraBd", 25) #choose your font type and font size
        canvas.drawString(40, 11*72, "Wewantwind Turbine Siting Report")
        canvas.setFont("OpenSansLt", 23) #choose your font type and font size
        canvas.drawString(40, 10.5*72, "Planning and other constraints")
        canvas.setFont("OpenSansExtraBd", 15) #choose your font type and font size
        canvas.drawString(40, 9.8*72, constraint['heading'])
        canvas.setFont("OpenSansLt", 11) #choose your font type and font size
        vpos = 9.5*72
        canvas.drawString(40, vpos, "Constraints shown in this section:")
        index = 0
        for dataset in constraint['datasets']:
            canvas.setFont("OpenSansLt", 11)
            canvas.setFillColorRGB(0,0,0)
            vpos -= 16
            if constraint == constraintslist[0]:
                canvas.setFont("OpenSansExtraBd", 11)
                color = constraint['layers'][0]['color']
                if len(constraint['layers']) > 0: color = constraint['layers'][index]['color']
                rgbarray = colors.to_rgb(color)
                canvas.setFillColorRGB(rgbarray[0], rgbarray[1], rgbarray[2])
            index += 1
            canvas.drawString(40, vpos, "• " + dataset)

        image = imagedirectory + '/' + constraint['heading'] + '.png'
        canvas.drawInlineImage(image, 40 , vpos - (48 * 7) - 100, width=72*7,height=420)
        canvas.setFont("OpenSansLt", 8)
        canvas.setFillColorRGB(0,0,0)
        canvas.drawString(40, 40, 'Constraints data from multiple sources and copyright of respective data providers - for full list refer to https://ckan.wewantwind.org')

        if constraint != lastcontraint:
            canvas.showPage()

    canvas.save()



id = str(uuid.uuid4())
lat = 51.12189892819
lng = -0.2345678966
# lat = 51
# lng = 0

filestem = str(lat) + "_" + str(lng)
downloadsdirectory = cwd + '/downloads/'
pdfpath = downloadsdirectory + filestem + '.pdf'
wordpath = downloadsdirectory + filestem + '.docx'
if True:
# if (os.path.isfile(wordpath) is False) or (os.path.isfile(pdfpath) is False):
    imagedirectory = processimages(id, [lng, lat], constraintslist, {'width': '600', 'height': '500', 'ratio': '3', 'zoom': '12', 'center': str(lng) + ',' + str(lat)})
    if os.path.isfile(wordpath) is False: createworddoc(wordpath, imagedirectory)
    if os.path.isfile(pdfpath) is False: createpdfdoc(pdfpath, imagedirectory)
    shutil.rmtree(imagedirectory)